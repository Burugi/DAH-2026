# -*- coding: utf-8 -*-
import pypandoc, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
md = r"C:\workspace\DAH2026_docs\코드팀_구현가이드.md"
dp = md[:-3] + ".docx"
pypandoc.convert_file(md, "docx", format="gfm", outputfile=dp)
doc = Document(dp); EAST = "Malgun Gothic"
def set_run(r, size=None, bold=None, color=None):
    r.font.name = EAST
    rpr = r._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = rpr.makeelement(qn('w:rFonts'), {}); rpr.insert(0, rf)
    for a in ('w:eastAsia','w:ascii','w:hAnsi'): rf.set(qn(a), EAST)
    if size is not None: r.font.size = Pt(size)
    if bold is not None: r.font.bold = bold
    if color is not None: r.font.color.rgb = color
def proc(p):
    st = (p.style.name if p.style else "") or ""; sz, bold, col = 10.5, None, None
    if st.startswith("Heading 1"): sz, bold, col = 15, True, RGBColor(0x1F,0x37,0x64)
    elif st.startswith("Heading 2"): sz, bold, col = 12.5, True, RGBColor(0x2E,0x59,0x84)
    elif st.startswith("Heading 3"): sz, bold = 11, True
    for r in p.runs: set_run(r, sz, bold, col)
    p.paragraph_format.line_spacing = 1.25
for p in doc.paragraphs: proc(p)
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs: set_run(r, 9.5)
nf = doc.styles['Normal'].font; nf.name = EAST; nf.size = Pt(10.5)
rpr = doc.styles['Normal'].element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
if rf is None: rf = rpr.makeelement(qn('w:rFonts'), {}); rpr.append(rf)
for a in ('w:eastAsia','w:ascii','w:hAnsi'): rf.set(qn(a), EAST)
doc.save(dp); print("DONE bytes:", os.path.getsize(dp))
